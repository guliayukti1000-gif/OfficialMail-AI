import io
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.shared import Pt, RGBColor


def build_pdf(subject: str, body: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=1 * inch, rightMargin=1 * inch,
        topMargin=1 * inch, bottomMargin=1 * inch,
    )
    styles = getSampleStyleSheet()
    subject_style = ParagraphStyle(
        "Subject", parent=styles["Heading1"], fontSize=16,
        textColor=colors.HexColor("#1E3A8A"), spaceAfter=18,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=11, leading=17,
        textColor=colors.HexColor("#0F172A"),
    )

    story = [Paragraph(subject, subject_style), Spacer(1, 12)]
    for para in body.split("\n\n"):
        clean = para.replace("\n", "<br/>")
        story.append(Paragraph(clean, body_style))
        story.append(Spacer(1, 10))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def build_docx(subject: str, body: str) -> bytes:
    doc = Document()
    heading = doc.add_heading(subject, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    for para in body.split("\n\n"):
        p = doc.add_paragraph(para)
        p.style.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
